"""
DOCX to MDP converter.

This module provides functions to convert DOCX (Word) documents to MDP format,
preserving structure and metadata.
"""

import os
import io
import datetime
from typing import Dict, Any, Union, Optional, List, Tuple
from pathlib import Path

try:
    import docx
    from docx.document import Document as DocxDocument
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

from mdp import Document, MDPFile
from mdp.metadata import create_metadata, generate_uuid, format_date
from .utils import normalize_metadata

def _check_docx_support():
    """
    Check if DOCX support is available.
    
    Raises:
        ImportError: If DOCX support is not available
    """
    if not DOCX_SUPPORT:
        raise ImportError(
            "DOCX support requires additional dependencies. "
            "Install with 'pip install datapack[converters]' or 'pip install python-docx'"
        )

def _extract_metadata_from_docx(doc: 'DocxDocument') -> Dict[str, Any]:
    """
    Extract metadata from a DOCX document.
    
    Args:
        doc: A python-docx Document object
        
    Returns:
        A dictionary of metadata
    """
    _check_docx_support()
    
    metadata = {}
    
    # Extract core properties
    core_props = doc.core_properties
    
    # Map DOCX properties to MDP metadata fields
    if core_props.title:
        metadata['title'] = core_props.title
    
    if core_props.author:
        metadata['author'] = core_props.author
    
    if core_props.created:
        metadata['created_at'] = format_date(core_props.created)
    
    if core_props.modified:
        metadata['updated_at'] = format_date(core_props.modified)
    
    if core_props.subject:
        metadata['description'] = core_props.subject
    
    if core_props.keywords:
        # Keywords are typically comma-separated
        metadata['tags'] = [tag.strip() for tag in core_props.keywords.split(',')]
    
    if core_props.version:
        metadata['version'] = core_props.version
    
    # Extract additional core properties as custom fields
    for prop_name in [
        'category', 'content_status', 'identifier', 'language',
        'last_modified_by', 'revision'
    ]:
        value = getattr(core_props, prop_name, None)
        if value:
            metadata[prop_name] = value
    
    # If no title found, try to extract from first heading
    if 'title' not in metadata and doc.paragraphs:
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading 1') and para.text.strip():
                metadata['title'] = para.text.strip()
                break
    
    # Normalize metadata to conform to MDP standards
    return normalize_metadata(metadata)

def _get_paragraph_style(paragraph: 'Paragraph') -> str:
    """
    Get the style of a paragraph.
    
    Args:
        paragraph: A python-docx Paragraph object
        
    Returns:
        The style name of the paragraph
    """
    _check_docx_support()
    
    style_name = paragraph.style.name.lower()
    
    # Map common Word styles to markdown
    if style_name.startswith('heading 1'):
        return 'h1'
    elif style_name.startswith('heading 2'):
        return 'h2'
    elif style_name.startswith('heading 3'):
        return 'h3'
    elif style_name.startswith('heading 4'):
        return 'h4'
    elif style_name.startswith('heading 5'):
        return 'h5'
    elif style_name.startswith('heading 6'):
        return 'h6'
    elif 'quote' in style_name or 'block quote' in style_name:
        return 'quote'
    elif 'list' in style_name and 'bullet' in style_name:
        return 'ul'
    elif 'list' in style_name and ('number' in style_name or 'decimal' in style_name):
        return 'ol'
    else:
        return 'p'

def _format_paragraph_text(paragraph: 'Paragraph') -> str:
    """
    Format the text of a paragraph with markdown formatting.
    
    Args:
        paragraph: A python-docx Paragraph object
        
    Returns:
        Markdown formatted text
    """
    _check_docx_support()
    
    text = paragraph.text
    
    # Apply formatting based on runs
    formatted_text = ""
    current_pos = 0
    
    for run in paragraph.runs:
        # Get the text of this run
        run_text = run.text
        
        # Skip empty runs
        if not run_text:
            continue
        
        # Find the position of this run in the paragraph text
        start_pos = text.find(run_text, current_pos)
        if start_pos == -1:
            # If not found, just append the text
            formatted_text += run_text
            continue
        
        # Add any text before this run
        if start_pos > current_pos:
            formatted_text += text[current_pos:start_pos]
        
        # Apply formatting
        if run.bold and run.italic:
            formatted_text += f"***{run_text}***"
        elif run.bold:
            formatted_text += f"**{run_text}**"
        elif run.italic:
            formatted_text += f"*{run_text}*"
        elif run.underline:
            formatted_text += f"__{run_text}__"
        elif run.strike:
            formatted_text += f"~~{run_text}~~"
        else:
            formatted_text += run_text
        
        # Update current position
        current_pos = start_pos + len(run_text)
    
    # Add any remaining text
    if current_pos < len(text):
        formatted_text += text[current_pos:]
    
    return formatted_text

def _table_to_markdown(table: 'Table') -> str:
    """
    Convert a DOCX table to markdown format.
    
    Args:
        table: A python-docx Table object
        
    Returns:
        Markdown representation of the table
    """
    _check_docx_support()
    
    if not table.rows:
        return ""
    
    result = []
    
    # Process header row
    header = []
    for cell in table.rows[0].cells:
        header.append(cell.text.strip() or " ")
    
    result.append("| " + " | ".join(header) + " |")
    
    # Add separator row
    result.append("| " + " | ".join(["---"] * len(header)) + " |")
    
    # Process data rows
    for row in table.rows[1:]:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip() or " ")
        result.append("| " + " | ".join(row_data) + " |")
    
    return "\n".join(result) + "\n\n"

def _docx_to_markdown(doc: 'DocxDocument') -> str:
    """
    Convert a DOCX document to markdown.
    
    Args:
        doc: A python-docx Document object
        
    Returns:
        Markdown representation of the document
    """
    _check_docx_support()
    
    content = []
    list_type = None
    list_count = 0
    
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            paragraph = Paragraph(element, doc)
            style = _get_paragraph_style(paragraph)
            text = _format_paragraph_text(paragraph)
            
            # Skip empty paragraphs
            if not text.strip():
                continue
            
            # Format based on style
            if style == 'h1':
                content.append(f"# {text}\n\n")
                list_type = None
            elif style == 'h2':
                content.append(f"## {text}\n\n")
                list_type = None
            elif style == 'h3':
                content.append(f"### {text}\n\n")
                list_type = None
            elif style == 'h4':
                content.append(f"#### {text}\n\n")
                list_type = None
            elif style == 'h5':
                content.append(f"##### {text}\n\n")
                list_type = None
            elif style == 'h6':
                content.append(f"###### {text}\n\n")
                list_type = None
            elif style == 'quote':
                content.append(f"> {text}\n\n")
                list_type = None
            elif style == 'ul':
                content.append(f"- {text}\n")
                list_type = 'ul'
            elif style == 'ol':
                list_count += 1
                content.append(f"{list_count}. {text}\n")
                list_type = 'ol'
            else:
                content.append(f"{text}\n\n")
                if list_type:
                    list_type = None
                    list_count = 0
        
        elif element.tag.endswith('tbl'):  # Table
            table = Table(element, doc)
            content.append(_table_to_markdown(table))
            list_type = None
            list_count = 0
    
    return ''.join(content)

def docx_to_mdp(
    docx_data: Union[str, Path, bytes, io.BytesIO], 
    output_path: Optional[Union[str, Path]] = None,
    metadata: Optional[Dict[str, Any]] = None,
    extract_metadata: bool = True
) -> Document:
    """
    Convert DOCX data to an MDP document.
    
    Args:
        docx_data: DOCX data as a file path, bytes, or BytesIO object
        output_path: Optional path to save the MDP file
        metadata: Optional metadata to use (overrides extracted metadata)
        extract_metadata: Whether to extract metadata from the DOCX
        
    Returns:
        An MDP Document object
        
    Raises:
        ImportError: If DOCX support is not available
        ValueError: If docx_data is invalid
    """
    _check_docx_support()
    
    # Load the DOCX data
    try:
        if isinstance(docx_data, (str, Path)):
            doc = docx.Document(docx_data)
            # Use filename as title if available
            filename = os.path.basename(docx_data)
            base_name = os.path.splitext(filename)[0]
            default_title = base_name.replace('_', ' ').replace('-', ' ').title()
        elif isinstance(docx_data, bytes):
            doc = docx.Document(io.BytesIO(docx_data))
            default_title = "Converted DOCX Document"
        elif isinstance(docx_data, io.BytesIO):
            doc = docx.Document(docx_data)
            default_title = "Converted DOCX Document"
        else:
            raise ValueError("Invalid DOCX data type. Expected file path, bytes, or BytesIO object.")
    except Exception as e:
        raise ValueError(f"Failed to read DOCX: {e}") from e
    
    # Extract metadata if requested
    doc_metadata = {}
    if extract_metadata:
        doc_metadata = _extract_metadata_from_docx(doc)
    
    # Override with provided metadata
    if metadata:
        # Normalize the provided metadata too
        normalized_metadata = normalize_metadata(metadata)
        doc_metadata.update(normalized_metadata)
    
    # Ensure required metadata
    if "title" not in doc_metadata:
        doc_metadata["title"] = default_title
    
    # Add source information
    doc_metadata["source"] = {
        "type": "docx",
        "converted_at": format_date(datetime.datetime.now()),
        "converter": "datapack.converters.docx_converter"
    }
    
    # Convert DOCX to markdown
    content = _docx_to_markdown(doc)
    
    # Create the document
    mdp_doc = Document.create(
        content=content,
        **doc_metadata
    )
    
    # Save if output path is provided
    if output_path:
        mdp_doc.save(output_path)
    
    return mdp_doc 